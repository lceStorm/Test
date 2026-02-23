# app_fixed9.py
# pip install streamlit python-docx lxml pillow
# (опционально, чтобы таймер обновлялся каждую секунду без кликов) pip install streamlit-autorefresh
# streamlit run app_fixed9.py
#
# ImageMagick (magick) нужен только чтобы конвертировать WMF/EMF -> PNG.
# Без него формулы-изображения могут не отображаться, но структура вопросов/ответов будет правильной.

import re
import json
import os
import tempfile
import subprocess
import hashlib
import shutil
import time
import random
import difflib
from typing import Dict, List, Optional, Tuple
from io import BytesIO

try:
    from PIL import Image
except Exception:
    Image = None

try:
    from streamlit_autorefresh import st_autorefresh  # type: ignore
except Exception:
    st_autorefresh = None  # fallback

import streamlit as st
from docx import Document
from docx.oxml.ns import qn

st.set_page_config(page_title="Тесты из DOCX/TXT", layout="wide", initial_sidebar_state="collapsed")



# --- UI: компактный режим (удобнее на телефоне) ---
if "compact_ui" not in st.session_state:
    st.session_state["compact_ui"] = True

with st.sidebar:
    st.markdown("### Интерфейс")
    st.session_state["compact_ui"] = st.checkbox("Компактный режим (телефон)", value=st.session_state["compact_ui"])

_COMPACT = bool(st.session_state.get("compact_ui", True))

if _COMPACT:
    st.markdown(
        """
        <style>
          /* уменьшаем верхний отступ и общий вертикальный “воздух” */
          .block-container {padding-top: 1.0rem; padding-bottom: 1.0rem;}
          /* делаем элементы чуть компактнее */
          [data-testid="stVerticalBlock"] {gap: 0.35rem;}
          /* уменьшаем отступы внутри экспандеров */
          details > summary {padding: 0.2rem 0;}
                  .qtitle{font-size:0.95rem;font-weight:600;line-height:1.2;margin:0.2rem 0 0.6rem 0;}
        </style>
        """,
        unsafe_allow_html=True
    )
# --- /UI ---

def safe_rerun():
    """Совместимость Streamlit: rerun/experimental_rerun."""
    try:
        st.rerun()
    except Exception:
        try:
            st.experimental_rerun()
        except Exception:
            pass


# -----------------------------
# Regex
# -----------------------------
OPTION_START_RE = re.compile(r"^\s*([A-EА-Е])\s*[\)\.\-–—:]\s*(.*?)(\s*\*)?\s*$")
QNUM_RE = re.compile(r"^\s*(\d{1,4})\s*[\.\)\-–—:]\s*.*$")
# Лояльный вариант: иногда в DOCX встречается "14 Какие ..." без точки после номера.
# Чтобы НЕ ловить строки-числа из вариантов ответов (например "42500"), требуем букву после пробела.
QNUM_LOOSE_RE = re.compile(r"^\s*(\d{1,4})\s+(?=[A-Za-zА-Яа-яЁё]).+$")
TICKET_RE = re.compile(r"^\s*Билет\s*№?\s*\d+\s*\.?\s*$", re.IGNORECASE)
IMG_TOKEN_RE = re.compile(r"\[IMG:(rId\d+)\]")


# Для устойчивого сопоставления между версиями DOCX:
# - В некоторых DOCX при пересохранении меняются байты картинок/формул → sha1 меняется.
# - Поэтому делаем «loose»-режим: считаем все картинки одинаковым токеном [IMG].
IMG_SIG_ANY_RE = re.compile(r"\[IMG:[^\]]+\]")

def _loosen_image_tokens(s: str) -> str:
    if not s:
        return ""
    return IMG_SIG_ANY_RE.sub("[IMG]", s)

NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "v": "urn:schemas-microsoft-com:vml",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
}


def split_inline_options(text: str) -> List[str]:
    """Разбивает строку, если в ней несколько вариантов подряд в ОДНОЙ строке:
       "А) ...; Б) ...; В) ...".

    Важно:
    - Не "ловит" инициалы внутри вариантов (например: "А) В. Майбах"), чтобы не
      превращать "В." в начало нового варианта.
    - В целом, если сомневается — лучше НЕ резать строку (это безопаснее для парсинга).
    """
    s = (text or "").strip()
    if not s:
        return [""]

    # Ищем возможные начала вариантов. Учитываем ) . - – — :
    cand: List[int] = []
    for m in re.finditer(r"(?:[A-EА-Е])\s*([\)\.\-–—:])", s):
        pos = m.start()
        delim = m.group(1)

        # Всегда допускаем вариант в начале строки
        if pos == 0:
            cand.append(pos)
            continue

        # Предыдущий НЕпробельный символ слева от кандидата
        j = pos - 1
        while j >= 0 and s[j] in " \t\u00A0":
            j -= 1
        prev = s[j] if j >= 0 else ""

        # 1) Частый кейс ошибки: после "А)" идёт "В. Фамилия" (инициал),
        #    где prev == ')'. Это НЕ начало нового варианта.
        if prev == ")":
            continue

        # 2) Для разделителя '.' (точка) режем только если слева реально разделитель,
        #    иначе это почти всегда инициалы/аббревиатуры внутри текста.
        if delim == "." and prev not in ";:\n\t-–—":
            continue

        cand.append(pos)

    # Если нашли 0 или 1 "начала" — резать нечего.
    if len(cand) <= 1:
        return [s]

    # Удаляем дубликаты и сортируем
    starts = sorted(set(cand))

    chunks: List[str] = []
    if starts[0] > 0:
        prefix = s[: starts[0]].strip().strip(";")
        if prefix:
            chunks.append(prefix)

    for i, st_pos in enumerate(starts):
        end = starts[i + 1] if i + 1 < len(starts) else len(s)
        part = s[st_pos:end].strip().strip(";")
        if part:
            chunks.append(part)

    return chunks



def paragraph_to_text_with_placeholders(paragraph) -> str:
    p = paragraph._p
    out: List[str] = []

    for child in p.iterchildren():
        tag = child.tag

        if tag == qn("w:r"):
            texts = child.findall(".//w:t", namespaces=NS)
            if texts:
                out.append("".join(t.text or "" for t in texts))

            blips = child.findall(".//a:blip", namespaces=NS)
            for blip in blips:
                rid = blip.get(qn("r:embed"))
                if rid:
                    out.append(f"[IMG:{rid}]")

            im_datas = child.findall(".//v:imagedata", namespaces=NS)
            for im in im_datas:
                rid = im.get(qn("r:id"))
                if rid:
                    out.append(f"[IMG:{rid}]")

        elif tag == qn("w:object"):
            im = child.find(".//v:imagedata", namespaces=NS)
            if im is not None:
                rid = im.get(qn("r:id"))
                if rid:
                    out.append(f"[IMG:{rid}]")
        else:
            im_datas = child.findall(".//v:imagedata", namespaces=NS)
            for im in im_datas:
                rid = im.get(qn("r:id"))
                if rid:
                    out.append(f"[IMG:{rid}]")

    text = "".join(out)
    text = re.sub(r"[ \t]+", " ", text).strip()
    return text


def _sha1(data: bytes) -> str:
    return hashlib.sha1(data).hexdigest()


def has_imagemagick() -> bool:
    """True if ImageMagick is available (v7: magick, v6: convert)."""
    return (shutil.which("magick") is not None) or (shutil.which("convert") is not None)

def imagemagick_cmd() -> list:
    """Return command prefix for conversion."""
    if shutil.which("magick") is not None:
        return ["magick", "convert"]
    return ["convert"]


def get_image_size_px(image_bytes: bytes) -> Optional[Tuple[int, int]]:
    if Image is None:
        return None
    try:
        with Image.open(BytesIO(image_bytes)) as im:
            return im.size
    except Exception:
        return None


def ensure_png_bytes(image_bytes: bytes, ext_hint: Optional[str] = None) -> Tuple[bytes, str, bool, str]:
    ext_hint = (ext_hint or "").lower().strip(".")
    needs_convert = ext_hint in {"wmf", "emf"}

    if not needs_convert and len(image_bytes) >= 4 and image_bytes[:4] == b"\xd7\xcd\xc6\x9a":
        needs_convert = True
        ext_hint = "wmf"

    if not needs_convert:
        return image_bytes, (ext_hint or "png"), True, ""

    if not has_imagemagick():
        return image_bytes, (ext_hint or "wmf"), False, "Нет ImageMagick (magick) — WMF/EMF не конвертируется."

    with tempfile.TemporaryDirectory() as td:
        in_path = os.path.join(td, f"in.{ext_hint or 'wmf'}")
        out_path = os.path.join(td, "out.png")
        with open(in_path, "wb") as f:
            f.write(image_bytes)

        try:
            run_kwargs = dict(check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            # На Windows ImageMagick может кратко открывать консольное окно — подавляем.
            if os.name == "nt":
                try:
                    run_kwargs["creationflags"] = subprocess.CREATE_NO_WINDOW
                    si = subprocess.STARTUPINFO()
                    si.dwFlags |= subprocess.STARTF_USESHOWWINDOW
                    run_kwargs["startupinfo"] = si
                except Exception:
                    pass
            subprocess.run(imagemagick_cmd() + [in_path, out_path],
                **run_kwargs,
            )
            with open(out_path, "rb") as f:
                return f.read(), "png", True, ""
        except Exception as e:
            return image_bytes, (ext_hint or "wmf"), False, f"Не удалось конвертировать WMF/EMF: {e}"


def extract_images_map(doc: Document) -> Dict[str, dict]:
    if "image_cache" not in st.session_state:
        st.session_state.image_cache = {}

    cache = st.session_state.image_cache
    images: Dict[str, dict] = {}

    for rid, rel in doc.part.rels.items():
        if "image" not in rel.reltype:
            continue

        image_part = rel._target
        blob = image_part.blob
        ext = (image_part.filename.split(".")[-1] if image_part.filename else "").lower()

        key = f"{rid}:{_sha1(blob)}"
        if key in cache:
            images[rid] = cache[key]
            continue

        out_bytes, out_ext, ok, note = ensure_png_bytes(blob, ext_hint=ext)
        dims = get_image_size_px(out_bytes)
        wpx, hpx = (dims if dims else (None, None))
        item = {"bytes": out_bytes, "ext": out_ext, "ok": ok, "note": note, "wpx": wpx, "hpx": hpx}
        cache[key] = item
        images[rid] = item

    return images


def collect_doc_lines(doc: Document) -> List[str]:
    lines: List[str] = []
    for p in doc.paragraphs:
        t = paragraph_to_text_with_placeholders(p)
        if t:
            lines.append(t)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = paragraph_to_text_with_placeholders(p)
                    if t:
                        lines.append(t)
    return lines


def parse_questions(lines: List[str]) -> List[dict]:
    questions: List[dict] = []

    q_lines: List[str] = []
    opts: Dict[str, str] = {}
    answer: Optional[str] = None

    in_options = False
    current_opt: Optional[str] = None
    current_ticket: Optional[str] = None

    def flush():
        nonlocal q_lines, opts, answer, in_options, current_opt
        if q_lines and opts:
            questions.append(
                {
                    "ticket": current_ticket,
                    "question": "\n".join(q_lines).strip(),
                    "options": dict(opts),
                    "answer": answer,
                }
            )
        q_lines[:] = []
        opts.clear()
        answer = None
        in_options = False
        current_opt = None

    for raw in lines:
        for chunk in split_inline_options(raw):
            line = (chunk or "").strip()
            if not line:
                continue

            if TICKET_RE.match(line):
                if q_lines or opts:
                    flush()
                current_ticket = re.sub(r"\s+", " ", line).strip().rstrip(".")
                continue

            if in_options and (QNUM_RE.match(line) or QNUM_LOOSE_RE.match(line)) and not OPTION_START_RE.match(line):
                flush()

            m = OPTION_START_RE.match(line)
            if m:
                in_options = True
                letter = m.group(1).upper()
                text = (m.group(2) or "").strip()
                star = (m.group(3) or "").strip()

                # лечение "Б)" дважды
                if letter == "Б" and "Б" in opts and "Г" not in opts:
                    letter = "Г"

                current_opt = letter
                if letter in opts and text:
                    opts[letter] = (opts[letter] + "\n" + text).strip()
                else:
                    opts[letter] = text

                if star:
                    answer = letter
                continue

            if not in_options:
                q_lines.append(line)
            else:
                if current_opt is None:
                    q_lines.append(line)
                else:
                    prev = opts.get(current_opt, "")
                    opts[current_opt] = (prev + "\n" + line).strip() if prev else line

    flush()
    return questions



# -----------------------------
# Parser for tagged format:
# <question>, <question2>, <question3> ... and <variant>
# -----------------------------
TAG_Q_RE = re.compile(r"^\s*<\s*question\d*\s*>\s*(.*)\s*$", re.IGNORECASE)
TAG_V_RE = re.compile(r"^\s*<\s*variant\s*>\s*(.*)\s*$", re.IGNORECASE)


def preprocess_tagged_lines(lines: List[str]) -> List[str]:
    """
    Подготавливает строки для тегированного формата (<question...>/<variant>),
    потому что в DOCX один абзац может содержать ВНУТРЕННИЕ переносы строки.
    Также иногда <variant> может встречаться внутри строки.

    На выходе — список строк, где каждый тег начинается с начала строки.
    """
    out: List[str] = []
    for ln in lines:
        if not ln:
            continue
        # 1) разрезаем абзац по внутренним переносам строки
        parts = re.split(r"[\r\n]+", str(ln))
        for part in parts:
            part = part.strip()
            if not part:
                continue

            # 2) если тег встретился в середине строки — делаем из него разделитель
            #    (вставляем перенос перед тегом, чтобы потом splitlines разнёс на отдельные строки)
            part = re.sub(r"(?i)\s*<\s*question(\d*)\s*>", r"\n<question\1>", part)
            part = re.sub(r"(?i)\s*<\s*variant\s*>", r"\n<variant>", part)

            for chunk in part.splitlines():
                chunk = chunk.strip()
                if chunk:
                    out.append(chunk)
    return out
RUS_LETTERS = ["А", "Б", "В", "Г", "Д", "Е", "Ж", "З", "И", "К"]  # запас

def parse_questions_tagged(lines: List[str]) -> List[dict]:
    """
    Формат:
      <question>Текст вопроса...
      <variant>Вариант 1
      <variant>Вариант 2
      ...
    Вопрос/вариант может продолжаться следующими строками без тега.
    """
    questions: List[dict] = []
    q_lines: List[str] = []
    variants: List[str] = []
    last_kind: Optional[str] = None  # "q" / "v"

    def flush():
        nonlocal q_lines, variants, last_kind
        q_text = "\n".join([x for x in q_lines if x.strip()]).strip()
        if q_text and variants:
            opts: Dict[str, str] = {}
            for i, v in enumerate(variants):
                key = RUS_LETTERS[i] if i < len(RUS_LETTERS) else f"V{i+1}"
                opts[key] = v.strip()
            questions.append({"question": q_text, "options": opts, "answer": None})
        q_lines = []
        variants = []
        last_kind = None

    have_seen_any_question = False

    for raw in lines:
        line = (raw or "").strip()
        if not line:
            continue

        m_q = TAG_Q_RE.match(line)
        if m_q:
            # новый вопрос => закрываем предыдущий
            if have_seen_any_question:
                flush()
            have_seen_any_question = True
            q_text = (m_q.group(1) or "").strip()
            q_lines = [q_text] if q_text else []
            variants = []
            last_kind = "q"
            continue

        m_v = TAG_V_RE.match(line)
        if m_v:
            # варианты до первого вопроса считаем мусором (в файле иногда встречаются)
            if not have_seen_any_question:
                continue
            v_text = (m_v.group(1) or "").strip()
            variants.append(v_text)
            last_kind = "v"
            continue

        # продолжение (без тега)
        if not have_seen_any_question:
            continue

        if last_kind == "v" and variants:
            variants[-1] = (variants[-1] + "\n" + line).strip()
        else:
            q_lines.append(line)
            last_kind = "q"

    flush()
    return questions


def parse_questions_auto(lines: List[str]) -> List[dict]:
    # авто-детект формата (теги или классика A)/Б)/...)
    # Важно: в DOCX теги могут быть внутри абзаца (с внутренними переносами строки),
    # поэтому детект делаем по подстроке, а для тег-формата прогоняем preprocess.
    for ln in lines[:500]:
        s = str(ln).lower()
        if "<question" in s or "<variant" in s:
            return parse_questions_tagged(preprocess_tagged_lines(lines))
    return parse_questions(lines)

def decode_text_bytes(b: bytes) -> str:
    for enc in ("utf-8-sig", "utf-8", "cp1251", "windows-1251", "latin-1"):
        try:
            return b.decode(enc)
        except Exception:
            continue
    return b.decode("utf-8", errors="replace")


def parse_docx_auto(file_like) -> Tuple[List[dict], Dict[str, dict]]:
    doc = Document(file_like)
    images_map = extract_images_map(doc)
    lines = collect_doc_lines(doc)
    questions = parse_questions_auto(lines)
    return questions, images_map


def parse_txt_auto(uploaded_file_or_bytes) -> Tuple[List[dict], Dict[str, dict]]:
    if isinstance(uploaded_file_or_bytes, (bytes, bytearray)):
        raw = bytes(uploaded_file_or_bytes)
    else:
        raw = uploaded_file_or_bytes.getvalue()

    text = decode_text_bytes(raw)
    lines: List[str] = []
    for ln in text.splitlines():
        ln = re.sub(r"[ \t]+", " ", ln).strip()
        if ln:
            lines.append(ln)
    questions = parse_questions_auto(lines)
    return questions, {}


# -----------------------------
# Rendering
# -----------------------------
def render_rich_text(s: str, images_map: Dict[str, dict]):
    last = 0
    any_img = False

    for m in IMG_TOKEN_RE.finditer(s or ""):
        any_img = True
        before = (s[last:m.start()] or "").strip()
        rid = m.group(1)
        last = m.end()

        if before:
            st.markdown(before)

        item = images_map.get(rid)
        if item and item.get("bytes"):
            if item.get("ok"):
                max_w = st.session_state.get("img_max_width", 500)
                show_ids = st.session_state.get("show_img_ids", False)
                orig_w = item.get("wpx")
                width = min(max_w, orig_w) if isinstance(orig_w, int) and orig_w > 0 else max_w
                if show_ids:
                    st.image(item["bytes"], width=width, caption=rid)
                else:
                    st.image(item["bytes"], width=width)
            else:
                st.info(f"Формула/картинка {rid} (нужно установить ImageMagick, чтобы показать).")
        else:
            st.warning(f"Не найдено изображение {rid}")

    tail = (s[last:] or "").strip()
    if tail:
        st.markdown(tail)

    if not any_img and not (s or "").strip():
        st.write("")


def render_rich_text_indented(s: str, images_map: Dict[str, dict], left_cols: int = 1, right_cols: int = 20):
    _, body = st.columns([left_cols, right_cols])
    with body:
        render_rich_text(s, images_map)


# -----------------------------
# Timer helpers
# -----------------------------
def fmt_mmss(seconds: int) -> str:
    seconds = max(0, int(seconds))
    m = seconds // 60
    s = seconds % 60
    return f"{m:02d}:{s:02d}"


def reset_timer(keep_settings: bool = True):
    st.session_state.timer_running = False
    st.session_state.timer_start_ts = None
    st.session_state.timer_end_ts = None
    st.session_state.timer_finish_reason = None
    st.session_state.timer_finish_ts = None
    if not keep_settings:
        st.session_state.timer_enabled = False
        st.session_state.timer_minutes = 20


def start_timer():
    minutes = int(st.session_state.timer_minutes or 0)
    if minutes <= 0:
        return
    now = time.time()
    st.session_state.timer_running = True
    st.session_state.timer_start_ts = now
    st.session_state.timer_end_ts = now + minutes * 60
    st.session_state.timer_finish_reason = None
    st.session_state.timer_finish_ts = None


def finish_to_results(reason: str = "manual"):
    # reason: manual | time
    st.session_state.test_phase = "results"
    if st.session_state.timer_running:
        st.session_state.timer_running = False
        st.session_state.timer_finish_reason = reason
        st.session_state.timer_finish_ts = time.time()


# -----------------------------
# Test helpers
# -----------------------------
def reset_testing_state():
    st.session_state.test_phase = "testing"  # testing | results | review
    st.session_state.test_index = 0
    st.session_state.review_pos = 0
    st.session_state.review_list = []
    st.session_state.user_answers = {}  # index -> letter
    st.session_state.answer_order_cache = {}
    # порядок вопросов (для перемешивания) пересоздаём заново
    st.session_state.test_order_indices = []
    st.session_state.test_order_sig = None
    reset_timer(keep_settings=True)


def compute_score(data: List[dict], user_answers: Dict[int, str], indices: List[int]) -> Tuple[int, int, float, List[int]]:
    total = len(indices)
    correct = 0
    wrong_indices: List[int] = []
    for i in indices:
        q = data[i]
        ans = user_answers.get(i)
        key = q.get("answer")
        if ans is not None and key is not None and ans == key:
            correct += 1
        else:
            wrong_indices.append(i)
    percent = (correct / total * 100.0) if total else 0.0
    return correct, total, percent, wrong_indices


def answered_count(indices: List[int], user_answers: Dict[int, str]) -> int:
    return sum(1 for i in indices if user_answers.get(i) is not None)


def get_test_indices(data: List[dict], only_marked: bool) -> List[int]:
    if only_marked:
        return [i for i, q in enumerate(data) if q.get("answer") is not None]
    return list(range(len(data)))


def _test_order_signature(file_sig: Optional[str], base_indices: List[int], only_marked: bool, shuffle: bool, seed: Optional[int]) -> str:
    # base_indices включаем в сигнатуру, чтобы при изменении набора вопросов порядок пересоздавался
    # (например, при переключении "только размеченные")
    return json.dumps(
        {
            "file": file_sig,
            "only_marked": bool(only_marked),
            "shuffle": bool(shuffle),
            "seed": int(seed) if seed is not None else None,
            "base": base_indices,
        },
        ensure_ascii=False,
        separators=(",", ":"),
    )


def prepare_test_order(base_indices: List[int]) -> List[int]:
    """
    Возвращает порядок прохождения вопросов для теста.
    Если включено "Перемешать вопросы", перемешивание стабильное по seed до перезапуска/новой перемешки.
    """
    file_sig = st.session_state.get("loaded_file_sig")
    only_marked = bool(st.session_state.get("test_only_marked"))
    shuffle = bool(st.session_state.get("shuffle_questions"))
    seed = st.session_state.get("shuffle_seed")

    if shuffle and seed is None:
        # один раз создаём seed (сохраняется в session_state)
        seed = int(time.time() * 1000) & 0x7fffffff
        st.session_state.shuffle_seed = seed

    sig = _test_order_signature(file_sig, base_indices, only_marked, shuffle, seed)

    # если сигнатура поменялась — пересоздаём порядок
    if (not st.session_state.get("test_order_indices")) or (st.session_state.get("test_order_sig") != sig):
        order = list(base_indices)
        if shuffle:
            rng = random.Random(seed)
            rng.shuffle(order)
        st.session_state.test_order_indices = order
        st.session_state.test_order_sig = sig

    return list(st.session_state.get("test_order_indices") or base_indices)


def apply_question_limit(order_indices: List[int]) -> List[int]:
    """Ограничивает количество вопросов в тесте (если включено в настройках).

    Берёт первые N вопросов из уже подготовленного порядка.
    Если включена перемешка вопросов, это фактически случайная выборка, фиксированная на тест.
    """
    if not bool(st.session_state.get("test_limit_enabled")):
        return list(order_indices)

    try:
        n = int(st.session_state.get("test_limit_count") or 0)
    except Exception:
        n = 0

    if n <= 0:
        return list(order_indices)

    if n >= len(order_indices):
        return list(order_indices)

    return list(order_indices[:n])



def prepare_option_order(global_idx: int, letters: List[str]) -> List[str]:
    """Возвращает порядок вариантов для конкретного вопроса в режиме тестирования.

    Если включено «Перемешать варианты ответов», перемешивание стабильное по seed.
    Опция «только при старте теста» фиксирует порядок один раз и хранит в session_state,
    чтобы он не менялся в процессе прохождения.
    """
    if not bool(st.session_state.get("shuffle_answers")):
        return list(letters)

    seed = st.session_state.get("shuffle_answers_seed")
    if seed is None:
        seed = int(time.time() * 1000) & 0x7fffffff
        st.session_state.shuffle_answers_seed = seed

    start_only = bool(st.session_state.get("shuffle_answers_start_only", True))
    if start_only:
        cache = st.session_state.get("answer_order_cache") or {}
        cached = cache.get(int(global_idx))
        if isinstance(cached, list) and len(cached) == len(letters) and set(cached) == set(letters):
            return list(cached)

    out = list(letters)
    # Отдельная стабильная "подперемешка" на каждый вопрос:
    rng = random.Random((int(seed) ^ (int(global_idx) * 1000003)) & 0x7fffffff)
    rng.shuffle(out)

    if start_only:
        cache = st.session_state.get("answer_order_cache") or {}
        cache[int(global_idx)] = list(out)
        st.session_state.answer_order_cache = cache

    return out

def _pick_display_letters(original_letters: List[str], n: int) -> List[str]:
    """Выбирает набор букв (A/B/C... или А/Б/В...) для отображения после перемешивания.

    Если в исходных буквах встречается кириллица — используем кириллицу.
    Иначе — латиницу. Длина всегда = n.
    """
    cyr = ["А","Б","В","Г","Д","Е","Ж","З","И","К","Л","М","Н","О","П","Р","С","Т","У","Ф","Х","Ц","Ч","Ш","Щ"]
    lat = ["A","B","C","D","E","F","G","H","I","J","K","L","M","N","O","P","Q","R","S","T","U","V","W","X","Y","Z"]

    # Определяем, есть ли кириллица в исходных метках
    has_cyr = any(bool(re.search(r"[\u0400-\u04FF]", str(x))) for x in original_letters)
    alphabet = cyr if has_cyr else lat

    if n <= len(alphabet):
        return alphabet[:n]

    # На всякий случай: если вариантов больше, чем букв, добавим индексы
    out = list(alphabet)
    k = 1
    while len(out) < n:
        out.append(f"{alphabet[-1]}{k}")
        k += 1
    return out[:n]


def prepare_option_view(global_idx: int, opts: Dict[str, str]) -> List[Tuple[str, str]]:
    """Возвращает список (display_letter, original_letter) для отображения вариантов в тестировании.

    - original_letter — ключ в opts (по нему проверяем правильность).
    - display_letter — буква, которую показываем пользователю.
      Если включена опция «перемешивать буквы вместе с ответами», display_letter переназначается как A/B/C...
    """
    ordered_orig = prepare_option_order(global_idx, list(opts.keys()))

    relabel = bool(st.session_state.get("shuffle_answers_relabel", True)) and bool(st.session_state.get("shuffle_answers"))
    if not relabel:
        return [(ol, ol) for ol in ordered_orig]

    display_letters = _pick_display_letters(ordered_orig, len(ordered_orig))
    return list(zip(display_letters, ordered_orig))



# -----------------------------
# Session init
# -----------------------------
defaults = {
    "data": None,
    "images_map": {},
    "mode": "Разметка ответов",
    "test_phase": "testing",
    "test_index": 0,
    "review_pos": 0,
    "review_list": [],
    "user_answers": {},
    "loaded_file_sig": None,
    "loaded_file_name": None,
    "autosave_key": True,
    "loaded_file_hash": None,
    "test_only_marked": True,
    "test_limit_enabled": False,
    "test_limit_count": 50,
    "shuffle_questions": False,
    "shuffle_seed": None,
    "shuffle_answers": False,
    "shuffle_answers_seed": None,
    "shuffle_answers_start_only": True,
    "shuffle_answers_relabel": True,
    "answer_order_cache": {},
    "test_order_indices": [],
    "test_order_sig": None,
    "mark_view_mode": "По одному (быстро)",
    "mark_index": 0,
    "mark_page": 0,
    "mark_page_size": 10,
    "mark_show_variants": True,
    "img_max_width": 500,
    "show_img_ids": False,
    "timer_enabled": False,
    "timer_minutes": 20,
    "timer_running": False,
    "timer_start_ts": None,
    "timer_end_ts": None,
    "timer_finish_reason": None,
    "timer_finish_ts": None,

    # UI: ручной импорт разметки из других версий документа
    "import_ui_open": False,

}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# -----------------------------
# Persist: сохраняем разметку (правильные ответы) между запусками
# -----------------------------
def _get_appdata_dir() -> str:
    base = os.getenv("LOCALAPPDATA") or os.getenv("APPDATA") or os.path.expanduser("~")
    path = os.path.join(base, "QuizTester")
    os.makedirs(path, exist_ok=True)
    return path


def _get_key_path(file_hash: str) -> str:
    keys_dir = os.path.join(_get_appdata_dir(), "answer_keys")
    os.makedirs(keys_dir, exist_ok=True)
    return os.path.join(keys_dir, f"{file_hash}.json")


def _extract_answer_list(data: List[dict]) -> List[Optional[str]]:
    answers: List[Optional[str]] = []
    for q in data:
        opts = q.get("options") or {}
        a = q.get("answer")
        answers.append(a if (a is not None and a in opts) else None)
    return answers


def save_persisted_key():
    """Сохраняет текущую разметку в AppData (если известен хеш файла)."""
    file_hash = st.session_state.get("loaded_file_hash")
    data = st.session_state.get("data")
    if not file_hash or not isinstance(data, list):
        return
    payload = {
        "file_hash": file_hash,
        "file_name": st.session_state.get("loaded_file_name"),
        "saved_at": int(time.time()),
        "answers": _extract_answer_list(data),
    }
    try:
        path = _get_key_path(str(file_hash))
        with open(path, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False)
        st.session_state.persist_last_saved_path = path
    except Exception:
        # Не валим приложение из-за сохранения
        pass


def load_persisted_key_into(data: List[dict], file_hash: str) -> int:
    """Загружает сохранённую разметку и применяет к data. Возвращает кол-во применённых ответов."""
    try:
        path = _get_key_path(str(file_hash))
        if not os.path.exists(path):
            return 0
        with open(path, "r", encoding="utf-8") as f:
            payload = json.load(f)
        answers = payload.get("answers")
        if not isinstance(answers, list):
            return 0

        applied = 0
        for i in range(min(len(data), len(answers))):
            a = answers[i]
            opts = data[i].get("options") or {}
            if a is not None and a in opts:
                data[i]["answer"] = a
                applied += 1
        st.session_state.persist_last_saved_path = path
        return applied
    except Exception:
        return 0


def delete_persisted_key(file_hash: Optional[str]) -> bool:
    try:
        if not file_hash:
            return False
        path = _get_key_path(str(file_hash))
        if os.path.exists(path):
            os.remove(path)
        return True
    except Exception:
        return False


# -----------------------------
# Persistence upgrade (v2): сохраняем разметку между версиями ДОКУМЕНТА
# Идея:
#  - раньше разметка привязывалась к хешу файла => любое сохранение DOCX меняет хеш и разметка "пропадает"
#  - теперь при отсутствии точного хеша мы пытаемся найти наиболее похожую сохранённую разметку и применить её
#  - в новых сохранениях храним соответствие "отпечаток вопроса -> отпечаток текста правильного варианта"
#    (устойчиво к перестановке вариантов/букв и мелким правкам форматирования)
# -----------------------------

_PERSIST_VERSION = 2

def _sha1_hex(b: bytes) -> str:
    return hashlib.sha1(b).hexdigest()

def _img_token_to_sig(match, images_map: Dict[str, dict]) -> str:
    """Заменяет [IMG:rIdX] на стабильный токен вида [IMG:sha1] (если есть байты картинки)."""
    rid = match.group(1)
    item = (images_map or {}).get(rid) or {}
    blob = item.get("bytes")
    if isinstance(blob, (bytes, bytearray)):
        return f"[IMG:{_sha1_hex(bytes(blob))[:12]}]"
    return "[IMG:missing]"

def _normalize_text_for_fp(s: str, images_map: Optional[Dict[str, dict]] = None) -> str:
    if not s:
        return ""
    # Привести токены картинок к стабильным (не зависящим от rId)
    if images_map is not None:
        s = IMG_TOKEN_RE.sub(lambda m: _img_token_to_sig(m, images_map), s)

    s = s.replace("\u00a0", " ").replace("\u200b", " ").replace("\ufeff", " ")
    s = s.replace("ё", "е").replace("Ё", "Е")
    s = re.sub(r"[ \t\r\n]+", " ", s).strip().lower()

    # Уберём служебные теги из "тегированного" DOCX, если вдруг остались в тексте
    s = re.sub(r"<\s*variant\s*>", " ", s, flags=re.IGNORECASE)
    s = re.sub(r"<\s*question\d*\s*>", " ", s, flags=re.IGNORECASE)

    # Нормализация пунктуации (чтобы мелкие различия не ломали совпадение)
    s = re.sub(r"[\"'`“”«»]+", "", s)
    s = re.sub(r"\s*[\.,;:!?()\[\]{}]+", " ", s)
    s = re.sub(r"[ \t\r\n]+", " ", s).strip()
    return s


def _normalize_text_for_fp_loose(s: str, images_map: Optional[Dict[str, dict]] = None) -> str:
    """
    Лояльная нормализация для отпечатка вопроса.
    Отличие от _normalize_text_for_fp:
      • любые картинки/формулы считаются одинаковыми ([IMG]),
        чтобы пересохранение DOCX не "ломало" сопоставление.
    """
    if not s:
        return ""
    # Любая картинка/формула → общий маркер
    s = IMG_TOKEN_RE.sub("[IMG]", s)

    s = s.replace("\u00a0", " ").replace("\u200b", " ").replace("\ufeff", " ")
    s = s.replace("ё", "е").replace("Ё", "Е")
    s = re.sub(r"[ \t\r\n]+", " ", s).strip().lower()

    # Уберём служебные теги из "тегированного" DOCX, если вдруг остались
    s = re.sub(r"<\s*variant\s*>", " ", s, flags=re.IGNORECASE)
    s = re.sub(r"<\s*question\d*\s*>", " ", s, flags=re.IGNORECASE)

    # Нормализация пунктуации
    s = re.sub(r"[\"'`“”«»]+", "", s)
    s = re.sub(r"\s*[\.,;:!?()\[\]{}]+", " ", s)
    s = re.sub(r"[ \t\r\n]+", " ", s).strip()
    return s

def _question_fingerprint(q: dict, images_map: Optional[Dict[str, dict]] = None) -> str:
    ticket = _normalize_text_for_fp(str(q.get("ticket") or ""), images_map)
    qtext = _normalize_text_for_fp(str(q.get("question") or ""), images_map)

    opts = q.get("options") or {}
    opt_norms = []
    for _k, v in (opts.items() if isinstance(opts, dict) else []):
        opt_norms.append(_normalize_text_for_fp(str(v), images_map))
    opt_norms = sorted([x for x in opt_norms if x])  # порядок вариантов не важен
    base = f"{ticket}|{qtext}|{'|'.join(opt_norms)}"
    return _sha1_hex(base.encode("utf-8"))


def _question_fingerprint_loose(q: dict, images_map: Optional[Dict[str, dict]] = None) -> str:
    """Отпечаток вопроса, устойчивый к пересохранению картинок/формул (все картинки = [IMG])."""
    ticket = _normalize_text_for_fp_loose(str(q.get("ticket") or ""), images_map)
    qtext = _normalize_text_for_fp_loose(str(q.get("question") or ""), images_map)

    opts = q.get("options") or {}
    opt_norms = []
    for _k, v in (opts.items() if isinstance(opts, dict) else []):
        opt_norms.append(_normalize_text_for_fp_loose(str(v), images_map))
    opt_norms = sorted([x for x in opt_norms if x])
    base = f"{ticket}|{qtext}|{'|'.join(opt_norms)}"
    return _sha1_hex(base.encode("utf-8"))

def _correct_option_fingerprint(q: dict, images_map: Optional[Dict[str, dict]] = None) -> Optional[str]:
    """Возвращает нормализованный текст правильного варианта (как отпечаток), либо None."""
    ans = q.get("answer")
    opts = q.get("options") or {}
    if not ans or not isinstance(opts, dict) or ans not in opts:
        return None
    return _normalize_text_for_fp(str(opts.get(ans, "")), images_map)



def _correct_option_fingerprint_loose(q: dict, images_map: Optional[Dict[str, dict]] = None) -> Optional[str]:
    ans = q.get("answer")
    opts = q.get("options") or {}
    if not ans or not isinstance(opts, dict) or ans not in opts:
        return None
    return _normalize_text_for_fp_loose(str(opts.get(ans, "")), images_map)

def _build_answer_text_map_loose(data: List[dict], images_map: Optional[Dict[str, dict]] = None) -> Dict[str, str]:
    """Карта: loose_fingerprint -> нормализованный текст правильного варианта (loose)."""
    m: Dict[str, str] = {}
    for q in data:
        fp = _question_fingerprint_loose(q, images_map)
        corr = _correct_option_fingerprint_loose(q, images_map)
        if corr:
            m[fp] = corr
    return m

def _build_qa_entries_loose(data: List[dict], images_map: Optional[Dict[str, dict]] = None) -> List[dict]:
    """Список записей (ticket/question/options/correct) для «умной» миграции при редактировании документа."""
    out: List[dict] = []
    for q in data:
        corr = _correct_option_fingerprint_loose(q, images_map)
        if not corr:
            continue
        ticket = _normalize_text_for_fp_loose(str(q.get("ticket") or ""), images_map)
        qtext = _normalize_text_for_fp_loose(str(q.get("question") or ""), images_map)
        opts = q.get("options") or {}
        opt_norms = []
        for _k, v in (opts.items() if isinstance(opts, dict) else []):
            nv = _normalize_text_for_fp_loose(str(v), images_map)
            if nv:
                opt_norms.append(nv)
        opt_norms = sorted(set(opt_norms))
        out.append({"t": ticket, "q": qtext, "o": opt_norms, "c": corr})
    return out
def _build_answer_text_map(data: List[dict], images_map: Optional[Dict[str, dict]] = None) -> Dict[str, str]:
    m: Dict[str, str] = {}
    for q in data:
        fp = _question_fingerprint(q, images_map)
        corr = _correct_option_fingerprint(q, images_map)
        if corr:
            m[fp] = corr
    return m

def _structure_sig(data: List[dict], images_map: Optional[Dict[str, dict]] = None) -> str:
    fps = [_question_fingerprint(q, images_map) for q in data]
    return _sha1_hex("|".join(fps).encode("utf-8"))

def _keys_dir() -> str:
    d = os.path.join(_get_appdata_dir(), "answer_keys")
    os.makedirs(d, exist_ok=True)
    return d

def _iter_saved_payloads() -> List[Tuple[str, dict]]:
    out: List[Tuple[str, dict]] = []
    try:
        for name in os.listdir(_keys_dir()):
            if not name.lower().endswith(".json"):
                continue
            path = os.path.join(_keys_dir(), name)
            try:
                with open(path, "r", encoding="utf-8") as f:
                    payload = json.load(f)
                if isinstance(payload, dict):
                    out.append((path, payload))
            except Exception:
                continue
    except Exception:
        pass
    return out

def _apply_payload_to_data(data: List[dict], payload: dict) -> int:
    """Применяет payload к data. Возвращает сколько ответов применено."""
    applied = 0
    images_map = st.session_state.get("images_map") or {}
    # Новый формат: answer_text_map по отпечаткам
    atm = payload.get("answer_text_map")
    if isinstance(atm, dict) and atm:
        for q in data:
            if q.get("answer") is not None:
                continue
            fp = _question_fingerprint(q, images_map)
            target = atm.get(fp)
            if not target:
                continue
            opts = q.get("options") or {}
            if not isinstance(opts, dict):
                continue
            for letter, text in opts.items():
                if _normalize_text_for_fp(str(text), images_map) == target:
                    q["answer"] = letter
                    applied += 1
                    break

    # 2) Loose формат: отпечатки без зависимости от байтов картинок/формул
    atm_loose = payload.get("answer_text_map_loose")
    if isinstance(atm_loose, dict) and atm_loose:
        for q in data:
            if q.get("answer") is not None:
                continue
            fp = _question_fingerprint_loose(q, images_map)
            target = atm_loose.get(fp)
            if not target:
                continue
            opts = q.get("options") or {}
            if not isinstance(opts, dict):
                continue
            for letter, text in opts.items():
                if _normalize_text_for_fp_loose(str(text), images_map) == target:
                    q["answer"] = letter
                    applied += 1
                    break

    # 3) Умная миграция: по похожести текста вопроса + пересечению вариантов
    # (помогает, если в документе поменяли форматирование/переносы/картинки и отпечатки не совпали)
    entries = payload.get("qa_entries_loose")
    if isinstance(entries, list) and entries:
        # индекс по билету для ускорения
        by_ticket = {}
        for e in entries:
            if not isinstance(e, dict):
                continue
            t = e.get("t") or ""
            by_ticket.setdefault(t, []).append(e)

        for q in data:
            if q.get("answer") is not None:
                continue
            t = _normalize_text_for_fp_loose(str(q.get("ticket") or ""), images_map)
            qtext = _normalize_text_for_fp_loose(str(q.get("question") or ""), images_map)

            opts = q.get("options") or {}
            cur_opts = []
            if isinstance(opts, dict):
                for _k, v in opts.items():
                    nv = _normalize_text_for_fp_loose(str(v), images_map)
                    if nv:
                        cur_opts.append(nv)
            cur_set = set(cur_opts)

            candidates = by_ticket.get(t) or entries
            best = None
            best_score = 0.0
            second = 0.0

            for e in candidates:
                et = e.get("t") or ""
                eq = e.get("q") or ""
                eo = e.get("o") or []
                # Сходство текста вопроса
                ts = difflib.SequenceMatcher(None, qtext, eq).ratio() if qtext and eq else 0.0
                # Пересечение вариантов (если есть)
                eo_set = set(eo) if isinstance(eo, list) else set()
                inter = len(cur_set & eo_set)
                union = len(cur_set | eo_set) or 1
                os = inter / union
                score = 0.75 * ts + 0.25 * os

                if score > best_score:
                    second = best_score
                    best_score = score
                    best = e
                elif score > second:
                    second = score

            # Порог и разрыв со 2-м местом — чтобы не назначать ошибочно
            if best and best_score >= 0.88 and (best_score - second) >= 0.04:
                target = best.get("c")
                if target and isinstance(opts, dict):
                    # 1) точное совпадение по loose-нормализации
                    found = False
                    for letter, text in opts.items():
                        if _normalize_text_for_fp_loose(str(text), images_map) == target:
                            q["answer"] = letter
                            applied += 1
                            found = True
                            break
                    # 2) если не нашли (очень редкий случай) — выберем самый похожий вариант
                    if (not found) and target:
                        best_letter = None
                        best_ls = 0.0
                        for letter, text in opts.items():
                            cand = _normalize_text_for_fp_loose(str(text), images_map)
                            ls = difflib.SequenceMatcher(None, cand, target).ratio() if cand else 0.0
                            if ls > best_ls:
                                best_ls = ls
                                best_letter = letter
                        if best_letter and best_ls >= 0.92:
                            q["answer"] = best_letter
                            applied += 1

    # Не выходим: дальше возможен старый формат (по индексу)

    # Старый формат: ответы по индексу
    answers = payload.get("answers")
    if isinstance(answers, list):
        for i in range(min(len(data), len(answers))):
            if data[i].get("answer") is not None:
                continue
            a = answers[i]
            opts = data[i].get("options") or {}
            if a is not None and isinstance(opts, dict) and a in opts:
                data[i]["answer"] = a
                applied += 1
    return applied


def _choose_best_payload(data: List[dict]) -> Tuple[Optional[str], Optional[dict], int]:
    """Ищем наиболее подходящий сохранённый payload для текущего документа."""
    images_map = st.session_state.get("images_map") or {}
    current_sig = _structure_sig(data, images_map)
    current_len = len(data)
    current_name = st.session_state.get("loaded_file_name") or ""

    best_score = -1
    best_path: Optional[str] = None
    best_payload: Optional[dict] = None
    best_applied_est = 0

    # Для ускорения overlap
    cur_fp_set = set(_question_fingerprint(q, images_map) for q in data)
    cur_fp_set_loose = set(_question_fingerprint_loose(q, images_map) for q in data)

    for path, payload in _iter_saved_payloads():
        score = 0

        # Приоритет: совпадение "структуры" (стойко к сохранениям/форматированию)
        if payload.get("structure_sig") == current_sig:
            score += 20000

        # Приоритет: похожее имя файла
        if current_name and payload.get("file_name") == current_name:
            score += 5000

        est = 0

        # Новый формат: overlap по отпечаткам
        atm = payload.get("answer_text_map")
        if isinstance(atm, dict) and atm:
            overlap = sum(1 for fp in cur_fp_set if fp in atm)
            score += overlap
            est = overlap

            # Доп. сигнал: loose-overlap (устойчив к пересохранению картинок/формул)
            atm_loose = payload.get("answer_text_map_loose")
            if isinstance(atm_loose, dict) and atm_loose:
                overlap2 = sum(1 for fp in cur_fp_set_loose if fp in atm_loose)
                score += int(overlap2 * 0.5)  # слабее, чтобы не перетягивал
                est = max(est, overlap2)

        else:
            # Старый формат: если длина совпадает, есть шанс применить по индексу
            ans = payload.get("answers")
            if isinstance(ans, list) and len(ans) == current_len:
                score += 100  # слабый сигнал
                est = sum(1 for a in ans if a is not None)
            else:
                # Есть записи для умной миграции — тоже небольшой сигнал
                entries = payload.get("qa_entries_loose")
                if isinstance(entries, list) and entries:
                    score += 20
                    est = len(entries)

        if score > best_score:
            best_score = score
            best_path = path
            best_payload = payload
            best_applied_est = est

    if best_score <= 0:
        return None, None, 0
    return best_path, best_payload, best_applied_est

# --- Переопределяем сохранение/загрузку разметки ---

def save_persisted_key():
    """Сохраняет текущую разметку в AppData (по хешу файла), + сохраняет карту отпечатков для миграции."""
    file_hash = st.session_state.get("loaded_file_hash")
    data = st.session_state.get("data")
    if not file_hash or not isinstance(data, list):
        return

    images_map = st.session_state.get("images_map") or {}

    payload = {
        "version": _PERSIST_VERSION,
        "file_hash": str(file_hash),
        "file_name": st.session_state.get("loaded_file_name"),
        "saved_at": int(time.time()),
        "answers": _extract_answer_list(data),  # для обратной совместимости/быстрой загрузки
        "structure_sig": _structure_sig(data, images_map),
        "answer_text_map": _build_answer_text_map(data, images_map),
        "answer_text_map_loose": _build_answer_text_map_loose(data, images_map),
        "qa_entries_loose": _build_qa_entries_loose(data, images_map),
    }
    try:
        path = _get_key_path(str(file_hash))
        with open(path, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False)
        st.session_state.persist_last_saved_path = path
    except Exception:
        pass


def load_persisted_key_into(data: List[dict], file_hash: str) -> int:
    """
    Пытается загрузить разметку:
      1) Сначала по точному хешу файла (как раньше).
      2) Если нет — ищет лучшую сохранённую разметку по "структуре"/имени/совпадениям и применяет.
    Возвращает кол-во применённых ответов.
    """
    # 1) точное совпадение по хешу
    try:
        path = _get_key_path(str(file_hash))
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                payload = json.load(f)
            applied = _apply_payload_to_data(data, payload if isinstance(payload, dict) else {})
            st.session_state.persist_last_saved_path = path
            return applied
    except Exception:
        pass

    # 2) миграция/поиск лучшего совпадения
    try:
        best_path, best_payload, _ = _choose_best_payload(data)
        if best_payload is None:
            return 0
        applied = _apply_payload_to_data(data, best_payload)
        if applied:
            st.session_state.persist_last_saved_path = best_path
            st.session_state.persist_migrated_from = best_path
        return applied
    except Exception:
        return 0
with st.sidebar:
    st.header("Сохранение")
    st.session_state.autosave_key = st.checkbox(
        "Автосохранять разметку ответов",
        value=bool(st.session_state.autosave_key),
        help="Разметка (правильные ответы) будет автоматически сохраняться в AppData и подгружаться при следующем запуске."
    )
    if st.session_state.loaded_file_hash:
        st.caption(f"Файл разметки: `{_get_key_path(str(st.session_state.loaded_file_hash))}`")
        c_del, c_save = st.columns([1, 1])
        with c_del:
            if st.button("🗑️ Очистить сохранённую разметку"):
                if delete_persisted_key(st.session_state.loaded_file_hash):
                    # не трогаем текущую разметку, только файл
                    st.success("Сохранённая разметка удалена.")
                else:
                    st.warning("Не удалось удалить файл разметки.")
        with c_save:
            if st.button("💾 Сохранить сейчас"):
                save_persisted_key()
                st.success("Разметка сохранена.")

        # --- Ручной импорт разметки из другой версии документа ---
        if st.button("📥 Импортировать разметку из другой версии"):
            st.session_state.import_ui_open = True
            st.rerun()

        if st.session_state.import_ui_open:
            saved = _iter_saved_payloads()

            if not saved:
                st.info("Сохранённых разметок не найдено.")
            else:
                current_path = _get_key_path(str(st.session_state.loaded_file_hash))
                items = []
                for p, payload in saved:
                    try:
                        if os.path.abspath(p) == os.path.abspath(current_path):
                            continue
                    except Exception:
                        pass

                    ts = payload.get("saved_at")
                    if isinstance(ts, (int, float)) and ts > 0:
                        dt = time.strftime("%Y-%m-%d %H:%M", time.localtime(int(ts)))
                    else:
                        dt = "?"

                    fname = payload.get("file_name") or os.path.basename(p)

                    # оценка количества ответов
                    est = 0
                    atm = payload.get("answer_text_map")
                    if isinstance(atm, dict) and atm:
                        est = len(atm)
                    else:
                        ans = payload.get("answers")
                        if isinstance(ans, list):
                            est = sum(1 for a in ans if a is not None)

                    label = f"{dt} — {fname} ({est} ответов) — {os.path.basename(p)}"
                    items.append((label, p, payload))

                if not items:
                    st.info("Других сохранений разметки нет.")
                else:
                    # По умолчанию выбираем "лучшее совпадение"
                    default_index = 0
                    try:
                        best_path, _, _ = _choose_best_payload(st.session_state.get("data") or [])
                        if best_path:
                            for i, it in enumerate(items):
                                if os.path.abspath(it[1]) == os.path.abspath(best_path):
                                    default_index = i
                                    break
                    except Exception:
                        pass

                    labels = [it[0] for it in items]
                    label_to = {it[0]: it for it in items}

                    choice = st.selectbox(
                        "Выберите сохранение",
                        labels,
                        index=min(default_index, max(0, len(labels) - 1)),
                        key="import_choice_label",
                    )

                    col_i1, col_i2 = st.columns(2)
                    with col_i1:
                        if st.button("✅ Импортировать", key="do_import"):
                            _, _, payload = label_to[choice]
                            data_ref = st.session_state.get("data") or []
                            applied = _apply_payload_to_data(data_ref, payload)
                            # Сохраним в текущий документ, чтобы дальше подхватывалось автоматически
                            save_persisted_key()
                            st.success(f"Импортировано ответов: {applied}")
                            st.session_state.import_ui_open = False
                            st.rerun()
                    with col_i2:
                        if st.button("✖️ Закрыть", key="close_import"):
                            st.session_state.import_ui_open = False
                            st.rerun()

                    st.caption("Импорт не перезаписывает уже размеченные ответы — заполняет только пустые.")

    st.divider()

    st.header("Отображение")
    st.session_state.img_max_width = st.slider(
        "Макс. ширина картинок (px)",
        min_value=200,
        max_value=1200,
        value=int(st.session_state.img_max_width),
        step=10,
    )
    st.session_state.show_img_ids = st.checkbox("Показывать подпись (rId)", value=bool(st.session_state.show_img_ids))

    st.divider()
    st.header("Тест")
    # При изменении параметров теста сбрасываем прогресс (чтобы порядок/набор вопросов не ломал результаты)
    _prev_only_marked = bool(st.session_state.test_only_marked)
    _prev_shuffle = bool(st.session_state.shuffle_questions)
    _prev_shuffle_answers = bool(st.session_state.shuffle_answers)
    _prev_limit_enabled = bool(st.session_state.get('test_limit_enabled'))
    _prev_limit_count = int(st.session_state.get('test_limit_count') or 0)
    st.session_state.test_only_marked = st.checkbox(
        "Тестировать только размеченные вопросы",
        value=bool(st.session_state.test_only_marked),
        help="Если размечены не все вопросы, включите этот режим, чтобы тестировать только те, где задан правильный ответ.",
    )

    st.session_state.shuffle_questions = st.checkbox(
        "Перемешать вопросы",
        value=bool(st.session_state.shuffle_questions),
        help="Перемешивает порядок вопросов в режиме тестирования. Порядок фиксируется на весь тест.",
    )

    st.session_state.test_limit_enabled = st.checkbox(
        "Ограничить количество вопросов",
        value=bool(st.session_state.get("test_limit_enabled")),
        help="Позволяет провести тест по ограниченному количеству вопросов. "
             "Будут взяты первые N вопросов из текущего порядка теста (после перемешивания, если оно включено).",
    )
    st.session_state.test_limit_count = st.number_input(
        "Количество вопросов в тесте",
        min_value=1,
        max_value=5000,
        value=int(st.session_state.get("test_limit_count") or 50),
        step=1,
        disabled=not bool(st.session_state.get("test_limit_enabled")),
    )

    if st.session_state.shuffle_questions:
        if st.session_state.shuffle_seed is None:
            st.session_state.shuffle_seed = int(time.time() * 1000) & 0x7fffffff
        col_s1, col_s2 = st.columns(2)
        with col_s1:
            if st.button("🔀 Новая перемешка"):
                st.session_state.shuffle_seed = int(time.time() * 1000) & 0x7fffffff
                st.session_state.answer_order_cache = {}
                reset_testing_state()
                safe_rerun()
        with col_s2:
            st.caption(f"Seed: {st.session_state.shuffle_seed}")
    else:
        st.session_state.shuffle_seed = None

    st.session_state.shuffle_answers = st.checkbox(
        "Перемешать варианты ответов",
        value=bool(st.session_state.shuffle_answers),
        help="Перемешивает порядок вариантов (A/Б/…) внутри каждого вопроса в режиме тестирования. Порядок фиксируется на весь тест.",
    )
    # Доп. опция: фиксировать порядок ответов только при старте теста
    if st.session_state.shuffle_answers:
        st.session_state.shuffle_answers_start_only = st.checkbox(
            "Перемешивать ответы только при старте теста",
            value=bool(st.session_state.shuffle_answers_start_only),
            help="Если включено, порядок вариантов фиксируется при старте/сбросе теста и не меняется в процессе прохождения.",
        )

        st.session_state.shuffle_answers_relabel = st.checkbox(
            "Перемешивать буквы вариантов вместе с ответами",
            value=bool(st.session_state.get("shuffle_answers_relabel", True)),
            help="Если включено, после перемешивания ответы перенумеровываются как A/B/C... (или А/Б/В...), чтобы не было 'перепутанных' букв.",
        )
    else:
        # когда перемешивание ответов выключено — сбрасываем кэш
        st.session_state.shuffle_answers_start_only = True
        st.session_state.answer_order_cache = {}

    if st.session_state.shuffle_answers:
        if st.session_state.shuffle_answers_seed is None:
            st.session_state.shuffle_answers_seed = int(time.time() * 1000) & 0x7fffffff
        col_a1, col_a2 = st.columns(2)
        with col_a1:
            started_answers = (st.session_state.get("mode") == "Тестирование") and (
                int(st.session_state.get("test_index") or 0) > 0 or len(st.session_state.get("user_answers") or {}) > 0
            )
            disable_shuffle_now = bool(st.session_state.get("shuffle_answers_start_only", True)) and started_answers
            if st.button("🔀 Новая перемешка ответов", disabled=disable_shuffle_now):
                st.session_state.shuffle_answers_seed = int(time.time() * 1000) & 0x7fffffff
                st.session_state.answer_order_cache = {}
                reset_testing_state()
                safe_rerun()
            if disable_shuffle_now:
                st.caption("Перемешка вариантов фиксируется на старте теста. Чтобы изменить порядок — нажмите «Начать заново» (тест начнётся заново).")
        with col_a2:
            st.caption(f"Seed: {st.session_state.shuffle_answers_seed}")
    else:
        st.session_state.shuffle_answers_seed = None

    # Если изменили параметры набора/порядка вопросов — сбросить тест
    if (bool(st.session_state.test_only_marked) != _prev_only_marked) or (bool(st.session_state.shuffle_questions) != _prev_shuffle) or (bool(st.session_state.shuffle_answers) != _prev_shuffle_answers or bool(st.session_state.get('test_limit_enabled')) != _prev_limit_enabled or int(st.session_state.get('test_limit_count') or 0) != _prev_limit_count):
        reset_testing_state()
        safe_rerun()

    st.session_state.timer_enabled = st.checkbox("Тестирование на время", value=bool(st.session_state.timer_enabled))
    st.session_state.timer_minutes = st.number_input(
        "Лимит времени (минут)",
        min_value=1,
        max_value=999,
        value=int(st.session_state.timer_minutes),
        step=1,
        disabled=not bool(st.session_state.timer_enabled),
    )

    col_t1, col_t2 = st.columns(2)
    with col_t1:
        if st.button("Старт таймера", disabled=not bool(st.session_state.timer_enabled) or bool(st.session_state.timer_running)):
            start_timer()
            safe_rerun()
    with col_t2:
        if st.button("Сброс таймера", disabled=not bool(st.session_state.timer_enabled) and not bool(st.session_state.timer_running)):
            reset_timer(keep_settings=True)
            safe_rerun()

if not has_imagemagick():
    st.warning(
        "ImageMagick (magick) не найден. WMF/EMF формулы могут не отображаться как картинки, "
        "но структура вопросов/ответов будет правильной."
    )

if st.session_state.timer_enabled and st_autorefresh is None:
    st.info("Чтобы таймер обновлялся автоматически каждую секунду: `pip install streamlit-autorefresh`")


uploaded = st.file_uploader("Загрузите .docx или .txt", type=["docx", "txt"])

if uploaded is not None:
    try:
        file_bytes = uploaded.getvalue()
        file_hash = _sha1(file_bytes)
        file_sig = f"{uploaded.name}:{file_hash}"

        if st.session_state.loaded_file_sig != file_sig:
            name = (uploaded.name or "").lower()
            if name.endswith(".txt"):
                parsed, images_map = parse_txt_auto(file_bytes)
            else:
                parsed, images_map = parse_docx_auto(BytesIO(file_bytes))

            if not parsed:
                st.error("Не удалось найти вопросы/варианты. Проверь формат: A) ... B) ... или А) ... Б) ...")
            else:
                st.session_state.data = parsed
                st.session_state.images_map = images_map
                st.session_state.loaded_file_sig = file_sig
                st.session_state.loaded_file_name = uploaded.name
                st.session_state.loaded_file_hash = file_hash
                reset_testing_state()

                # сброс разметки (ускоренная разметка по одному вопросу)
                st.session_state.mark_index = 0
                st.session_state.mark_page = 0
                st.session_state.mark_jump = 1
                st.session_state.mark_jump_dirty = True

                # если размечены не все — по умолчанию тестируем только размеченные
                st.session_state.test_only_marked = any(q.get("answer") is None for q in parsed)

                                # Автоподгрузка сохранённой разметки (если есть)
                applied = load_persisted_key_into(parsed, file_hash)
                if applied:
                    st.info(f"Загружена сохранённая разметка: {applied}/{len(parsed)}")
                st.success(f"Найдено вопросов: {len(parsed)}")

    except Exception as e:
        st.exception(e)

data: Optional[List[dict]] = st.session_state.data
images_map: Dict[str, dict] = st.session_state.images_map

if data is None:
    st.info("Загрузите файл, чтобы начать.")
    st.stop()

if st.session_state.loaded_file_name:
    st.caption(f"Файл: {st.session_state.loaded_file_name} · Вопросов: {len(data)}")

st.session_state.mode = st.radio("Режим", ["Разметка ответов", "Тестирование"], horizontal=True)

# =============================
# Разметка
# =============================
if st.session_state.mode == "Разметка ответов":
    st.subheader("Разметка ответов")

    total = len(data)

    # Если callback (автопереход) попросил перейти к другому вопросу — делаем это в начале рерана,
    # чтобы новый индекс применился без лишних rerun и без конфликтов с виджетами.
    _pending = st.session_state.pop("_mark_pending_go", None)
    if _pending is not None:
        try:
            st.session_state.mark_index = max(0, min(total - 1, int(_pending)))
        except Exception:
            st.session_state.mark_index = 0
        st.session_state.mark_jump_dirty = True

    # Прогресс разметки (считаем отмеченным, если ответ задан и он существует среди вариантов)
    marked = 0
    for q in data:
        opts = q.get("options") or {}
        if (q.get("answer") is not None) and (q.get("answer") in opts):
            marked += 1
    st.caption(f"Размечено: {marked}/{total}")

    # --- Верхняя панель управления (без expander) ---
    c0, c1, c2, c3, c4 = st.columns([1.2, 1.6, 1.6, 1.4, 2.2])

    with c0:
        if st.button("Сбросить ответы тестирования"):
            reset_testing_state()
            st.success("Ответы тестирования сброшены.")

    with c1:
        view_modes = ["По одному (быстро)", "Страницей", "Списком (медленно)"]
        if st.session_state.get("mark_view_mode") not in view_modes:
            st.session_state.mark_view_mode = view_modes[0]
        st.selectbox("Отображение", view_modes, key="mark_view_mode")

    with c2:
        st.checkbox("Показывать варианты", key="mark_show_variants")
        if "mark_auto_advance" not in st.session_state:
            st.session_state.mark_auto_advance = True
        st.checkbox("Автопереход к следующему неразмеченному", key="mark_auto_advance")

    with c3:
        st.number_input("Размер страницы", min_value=5, max_value=50, step=5, key="mark_page_size")

    with c4:
        # "прыжок" к вопросу (без лагов и без ошибки session_state)
        if "mark_jump_dirty" not in st.session_state:
            st.session_state.mark_jump_dirty = True

        # Синхронизируем число в инпуте ТОЛЬКО когда индекс менялся программно
        if ("mark_jump" not in st.session_state) or st.session_state.mark_jump_dirty:
            st.session_state.mark_jump = int(st.session_state.mark_index) + 1
            st.session_state.mark_jump_dirty = False

        st.number_input("Перейти к №", min_value=1, max_value=total, step=1, key="mark_jump")
        new_index = int(st.session_state.mark_jump) - 1
        if new_index != int(st.session_state.mark_index):
            st.session_state.mark_index = new_index

    def set_mark_index(i: int):
        i = max(0, min(total - 1, int(i)))
        st.session_state.mark_index = i
        st.session_state.mark_jump_dirty = True


    def find_next_unmarked(start_idx: int) -> Optional[int]:
        """Следующий вопрос без правильного ответа (циклически)."""
        # Ищем после текущего
        for j in range(start_idx + 1, total):
            if data[j].get("answer") is None and (data[j].get("options") or {}):
                return j
        # Если не нашли — идём с начала
        for j in range(0, start_idx + 1):
            if data[j].get("answer") is None and (data[j].get("options") or {}):
                return j
        return None

    def render_mark_item(i: int, show_header: bool = True):
        q = data[i]
        ticket = q.get("ticket")

        if show_header:
            header = f"### {i+1}."
            if ticket:
                header += f"  ·  {ticket}"
            st.markdown(header)

        render_rich_text(q.get("question", ""), images_map)

        opts = q.get("options") or {}
        if not opts:
            st.warning("Нет вариантов у вопроса.")
            return

        letters = list(opts.keys())
        current = q.get("answer") if q.get("answer") in letters else letters[0]

        # Callback: сохраняем выбранный правильный ответ сразу
        def _commit():
            st.session_state.data[i]["answer"] = st.session_state[f"mark_choice_{i}"]
            # Автосохранение разметки
            if bool(st.session_state.autosave_key):
                save_persisted_key()

            # Автопереход: после выбора (кроме подтверждения A по умолчанию) прыгаем на следующий неразмеченный.
            if bool(st.session_state.get("mark_auto_advance")) and st.session_state.get("mark_view_mode") == "По одному (быстро)":
                nxt = find_next_unmarked(i)
                if nxt is not None and nxt != i:
                    # Срабатывает в начале следующего рерана
                    st.session_state["_mark_pending_go"] = int(nxt)

        st.radio(
            "Правильный ответ:",
            letters,
            index=letters.index(current),
            key=f"mark_choice_{i}",
            horizontal=True,
            on_change=_commit,
        )
        if q.get("answer") is None:
            if st.button("✅ Подтвердить текущий ответ", key=f"mark_confirm_{i}"):
                st.session_state.data[i]["answer"] = st.session_state.get(f"mark_choice_{i}", current)
                # Автосохранение разметки при подтверждении
                if bool(st.session_state.autosave_key):
                    save_persisted_key()
                if bool(st.session_state.get("mark_auto_advance")) and st.session_state.get("mark_view_mode") == "По одному (быстро)":
                    nxt = find_next_unmarked(i)
                    if nxt is not None and nxt != i:
                        set_mark_index(int(nxt))
                safe_rerun()


        if st.session_state.get("mark_show_variants", True):
            for L in letters:
                st.markdown(f"**{L})**")
                render_rich_text_indented(opts.get(L, ""), images_map)
        else:
            st.caption("Варианты скрыты для ускорения. Включите «Показывать варианты» при необходимости.")

    # --- Основной вывод ---
    view = st.session_state.get("mark_view_mode", "По одному (быстро)")

    if view == "По одному (быстро)":
        idx = int(st.session_state.mark_index)

        next_unmarked = find_next_unmarked(idx)

        nav1, nav2, nav3, nav4 = st.columns([1, 1.8, 1, 5])
        with nav1:
            if st.button("⬅️ Назад", disabled=(idx <= 0)):
                set_mark_index(idx - 1)
                safe_rerun()
        with nav2:
            if st.button("⏭️ Следующий неразмеченный", disabled=(next_unmarked is None)):
                set_mark_index(int(next_unmarked))
                safe_rerun()
        with nav3:
            if st.button("Вперёд ➡️", disabled=(idx >= total - 1)):
                set_mark_index(idx + 1)
                safe_rerun()
        with nav4:
            st.write(f"Вопрос **{idx+1}** из **{total}**")

        st.divider()
        render_mark_item(idx, show_header=True)
        st.divider()

    else:
        page_size = int(st.session_state.get("mark_page_size", 10) or 10)
        pages = (total + page_size - 1) // page_size

        # Если пользователь выбрал "Списком", показываем предупреждение
        if view == "Списком (медленно)":
            st.warning("Списком может подлагивать на больших файлах. Лучше «По одному» или «Страницей».")
            page = 0
            page_size = total
        else:
            # Страничный режим
            if "mark_page" not in st.session_state:
                st.session_state.mark_page = 0
            page = int(st.session_state.mark_page)
            page = max(0, min(pages - 1, page))

            p1, p2, p3, p4 = st.columns([1, 1, 2, 6])
            with p1:
                if st.button("⬅️ Стр.", disabled=(page <= 0)):
                    st.session_state.mark_page = page - 1
                    safe_rerun()
            with p2:
                if st.button("Стр. ➡️", disabled=(page >= pages - 1)):
                    st.session_state.mark_page = page + 1
                    safe_rerun()
            with p3:
                new_page = st.number_input("Страница", 1, pages, value=page + 1, step=1, key="mark_page_input")
                if int(new_page) - 1 != page:
                    st.session_state.mark_page = int(new_page) - 1
                    safe_rerun()
            with p4:
                st.write(f"Показ: **{page+1}/{pages}**  ·  по **{page_size}** вопросов")

        start_i = page * page_size
        end_i = min(total, start_i + page_size)

        for i in range(start_i, end_i):
            render_mark_item(i, show_header=True)
            st.divider()

    # --- Экспорт ключа ---
    key_json = json.dumps(
        [{"question": q["question"], "options": q["options"], "answer": q.get("answer")} for q in data],
        ensure_ascii=False,
        indent=2,
    )
    st.download_button(
        "Скачать ключ ответов (JSON)",
        key_json.encode("utf-8"),
        "answer_key.json",
        "application/json",
    )

# =============================
# Тестирование
# =============================
else:
    st.subheader("Тестирование")

    c1, c2, c3 = st.columns([1, 1, 2])
    with c1:
        if st.button("Начать заново"):
            reset_testing_state()
            st.success("Тест начат заново.")
    with c2:
        if st.button("К результатам"):
            finish_to_results(reason="manual")
            safe_rerun()

    # База теста (все или только размеченные)
    base_indices = get_test_indices(data, bool(st.session_state.test_only_marked))

    if not base_indices:
        st.warning("Пока нет ни одного размеченного вопроса. Перейдите в «Разметка ответов» и отметьте хотя бы один вопрос.")
        st.stop()

    # Если тестируем ВСЕ вопросы, но размечены не все — нельзя
    if (not bool(st.session_state.test_only_marked)) and any(q.get("answer") is None for q in data):
        st.warning(
            "Есть вопросы без заданного правильного ответа. Перейдите в «Разметка ответов» и размечайте дальше "
            "или включите «Тестировать только размеченные вопросы» в боковой панели."
        )
        st.stop()

        # Порядок вопросов для теста (с перемешиванием)
    test_indices = apply_question_limit(prepare_test_order(base_indices))

# -----------------------------
    # Timer (только для основного теста, не для работы над ошибками)
    # -----------------------------
    if st.session_state.test_phase == "testing" and bool(st.session_state.timer_enabled):
        if not bool(st.session_state.timer_running):
            st.warning("Таймер включён, но не запущен. Нажмите «Старт таймера» в боковой панели (или начните тест без таймера).")
        else:
            # авто-обновление раз в секунду (если установлен streamlit-autorefresh)
            if st_autorefresh is not None:
                st_autorefresh(interval=1000, key="timer_tick")

            now = time.time()
            remaining = int((st.session_state.timer_end_ts or now) - now)
            total_seconds = int((st.session_state.timer_minutes or 0) * 60)
            used = total_seconds - max(0, remaining)

            tcol1, tcol2, tcol3 = st.columns([1, 1, 2])
            with tcol1:
                st.metric("Осталось", fmt_mmss(remaining))
            with tcol2:
                st.metric("Прошло", fmt_mmss(used))
            with tcol3:
                st.info("Тест можно завершить в любой момент. Неотвеченные будут засчитаны как неверные.")

            if remaining <= 0:
                finish_to_results(reason="time")
                safe_rerun()

    # -----------------------------
    # RESULTS
    # -----------------------------
    if st.session_state.test_phase == "results":
        correct_n, total_n, percent, wrong_indices = compute_score(data, st.session_state.user_answers, test_indices)

        st.markdown("## Результат")
        m1, m2, m3, m4 = st.columns(4)
        
        with m1:
            st.metric("Процент правильных", f"{percent:.1f}%")
        with m2:
            st.metric("Правильных", f"{correct_n}/{total_n}")

        answered_n = answered_count(test_indices, st.session_state.user_answers)
        incorrect_n = 0
        for _i in test_indices:
            ua = st.session_state.user_answers.get(_i)
            key = data[_i].get("answer")
            if ua is not None and (key is None or ua != key):
                incorrect_n += 1
        skipped_n = max(0, total_n - answered_n)

        with m3:
            st.metric("Неправильных", f"{incorrect_n}")
        with m4:
            st.metric("Отвечено", f"{answered_n}/{total_n}")
        if skipped_n:
            st.caption(f"Пропущено: {skipped_n}")

# таймерная сводка
        if bool(st.session_state.timer_enabled):
            minutes = int(st.session_state.timer_minutes or 0)
            limit = minutes * 60
            if st.session_state.timer_start_ts is not None:
                end_ts = st.session_state.timer_finish_ts or time.time()
                elapsed = int(max(0, min(limit, end_ts - st.session_state.timer_start_ts)))
                reason = st.session_state.timer_finish_reason
                label = "Время вышло" if reason == "time" else ("Завершено вручную" if reason == "manual" else "—")
                st.info(f"Таймер: лимит {fmt_mmss(limit)} · прошло {fmt_mmss(elapsed)} · {label}")

        if wrong_indices:
            st.warning(f"Неправильных/неотвеченных: {len(wrong_indices)}")

            show_errors = st.checkbox("Показать ошибки (вопросы с неверным/пустым ответом)", value=False)
            if show_errors:
                for idx in wrong_indices:
                    q = data[idx]
                    ticket = q.get("ticket")
                    st.markdown(f"### Вопрос {idx+1}" + (f" · {ticket}" if ticket else ""))
                    render_rich_text(q["question"], images_map)
                    your = st.session_state.user_answers.get(idx)
                    st.write(f"Ваш ответ: {your if your is not None else '—'}")
                    st.write(f"Правильный ответ: {q.get('answer')}")
                    st.divider()

            colx, coly = st.columns([1, 2])
            with colx:
                if st.button("Работа над ошибками"):
                    st.session_state.review_list = list(wrong_indices)
                    st.session_state.review_pos = 0
                    for i in st.session_state.review_list:
                        st.session_state.user_answers.pop(i, None)
                    st.session_state.test_phase = "review"
                    st.success("Открыта работа над ошибками.")
                    safe_rerun()
            with coly:
                st.info("В режиме «Работа над ошибками» вы заново отвечаете только на неправильные вопросы (без таймера).")

        else:
            st.success("Все ответы верны! 🎉")

        st.stop()

    # -----------------------------
    # TESTING / REVIEW renderer
    # -----------------------------
    is_review = st.session_state.test_phase == "review"
    if is_review and not st.session_state.review_list:
        st.info("Нет ошибок для работы над ними. Перейдите «К результатам» или начните заново.")
        st.stop()

    order_indices = st.session_state.review_list if is_review else test_indices
    pos_key = "review_pos" if is_review else "test_index"
    pos = st.session_state[pos_key]
    pos = max(0, min(pos, len(order_indices) - 1))
    st.session_state[pos_key] = pos

    global_idx = order_indices[pos]
    q = data[global_idx]
    opts = q["options"]
    correct = q.get("answer")

    answered = answered_count(order_indices, st.session_state.user_answers)
    total_phase = len(order_indices)

    st.progress(answered / total_phase if total_phase else 0.0)
    st.caption(
        ("Работа над ошибками: " if is_review else "Тест: ")
        + f"ответов {answered}/{total_phase} · всего вопросов {len(data)}"
    )

    def go_prev():
        st.session_state[pos_key] = max(0, st.session_state[pos_key] - 1)

    def go_next():
        st.session_state[pos_key] = min(len(order_indices) - 1, st.session_state[pos_key] + 1)

    nav1, nav2, nav3, nav4 = st.columns([1, 1, 1, 2])
    with nav1:
        st.button("⬅️ Назад", on_click=go_prev, disabled=(pos == 0), key=f"prev_{st.session_state.test_phase}")
    with nav2:
        st.button(
            "Вперёд ➡️",
            on_click=go_next,
            disabled=(pos == len(order_indices) - 1),
            key=f"next_{st.session_state.test_phase}",
        )
    with nav3:
        # ✅ досрочное завершение — всегда доступно
        finish_label = "Завершить работу" if is_review else "Завершить тест"
        if st.button(f"{finish_label} сейчас", key=f"finish_any_{st.session_state.test_phase}"):
            finish_to_results(reason="manual")
            safe_rerun()
    with nav4:
        st.info("Можно завершить в любой момент — неотвеченные будут считаться неверными.")

    base_title = f"Вопрос {global_idx+1} из {len(data)}"
if is_review:
    base_title += f"  ·  Ошибка {pos+1} из {len(order_indices)}"
ticket = q.get("ticket")
if ticket:
    base_title += f"  ·  {ticket}"

if _COMPACT:
    st.markdown(f"<div class='qtitle'>{base_title}</div>", unsafe_allow_html=True)
else:
    st.markdown("## " + base_title)

    render_rich_text(q.get("question", ""), images_map)

    if not opts:
        st.warning("У этого вопроса нет вариантов.")
        st.stop()
    if correct is None:
        st.warning("Не задан правильный ответ. Перейдите в “Разметка ответов” или включите «Тестировать только размеченные вопросы».")
        st.stop()

    selected = st.session_state.user_answers.get(global_idx)
    st.markdown("### Варианты (нажмите на букву, чтобы выбрать)")
    view = prepare_option_view(global_idx, opts)

    # Защита от некорректного формата view (иногда в окружениях/после правок может прийти не список пар).
    # Нормализуем к виду List[Tuple[display, original]].
    norm_view = []
    try:
        for item in list(view):
            if isinstance(item, (list, tuple)) and len(item) == 2:
                disp, orig = item
            elif isinstance(item, str):
                disp = orig = item
            else:
                continue
            norm_view.append((str(disp), str(orig)))
    except Exception:
        norm_view = []

    if not norm_view:
        # Фоллбэк: показываем как есть, без переназначения букв
        norm_view = [(k, k) for k in prepare_option_order(global_idx, list(opts.keys()))]

    view = norm_view
    display_by_orig = {orig: disp for (disp, orig) in view}
    orig_by_display = {disp: orig for (disp, orig) in view}

    for disp, orig in view:
        # Более устойчиво на телефоне: не индексируем columns (row[0]/row[1]),
        # а распаковываем. Если columns по какой-то причине недоступны, делаем фоллбэк без колонок.
        try:
            col_btn, col_txt = st.columns([1, 25])
        except Exception:
            col_btn = st.container()
            col_txt = st.container()

        with col_btn:
            if st.button(f"{disp})", key=f"pick_{st.session_state.test_phase}_{global_idx}_{orig}"):
                st.session_state.user_answers[global_idx] = orig
                selected = orig
            if selected == orig:
                st.markdown("✅")

        with col_txt:
            render_rich_text(opts[orig], images_map)
        st.write("")

    selected_disp = display_by_orig.get(selected, str(selected)) if selected is not None else None
    correct_disp = display_by_orig.get(correct, str(correct)) if correct is not None else None

    if selected is not None:
        if selected == correct:
            st.markdown(
                "<div style='padding:10px;border-radius:10px;background:#123b22;border:1px solid #1f8a3b;color:#ffffff;'>"
                f"✅ Верно (ваш ответ: {selected_disp})</div>",
                unsafe_allow_html=True,
            )
        else:
            st.markdown(
                "<div style='padding:10px;border-radius:10px;background:#3b1212;border:1px solid #c2303a;color:#ffffff;'>"
                f"❌ Неверно (ваш ответ: {selected_disp}). Правильный ответ: {correct_disp}</div>",
                unsafe_allow_html=True,
            )
    else:
        st.info("Вы ещё не ответили на этот вопрос.")


    def jump_next_unanswered():
        for j in range(pos + 1, len(order_indices)):
            gi = order_indices[j]
            if st.session_state.user_answers.get(gi) is None:
                st.session_state[pos_key] = j
                return
        for j in range(0, pos):
            gi = order_indices[j]
            if st.session_state.user_answers.get(gi) is None:
                st.session_state[pos_key] = j
                return

    if st.button("Перейти к следующему неотвеченному", key=f"jump_{st.session_state.test_phase}"):
        jump_next_unanswered()
        safe_rerun()

    st.caption("Зелёный — верно, красный — неверно. Результат будет показан после завершения (вручную или по таймеру).")
